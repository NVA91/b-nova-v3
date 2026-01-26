#!/usr/bin/env python3
"""
NOVA v3 - Telegram Voice Assistant
Sprach-zu-Text-zu-Dokument Pipeline mit Telegram Bot Integration

Features:
- Empfang von Sprachnachrichten via Telegram
- Transkription mit Whisper (lokal oder OpenAI API)
- Dokumenten-Generierung (Markdown, PDF, DOCX)
- Integration mit NOVA Backend für Persistenz
"""

import os
import logging
import asyncio
from pathlib import Path
from typing import Optional
import tempfile
from datetime import datetime

from telegram import Update
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    filters,
    ContextTypes,
)
import httpx
import whisper
from docx import Document
from markdown2 import markdown
from weasyprint import HTML

# Konfiguration
TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
WHISPER_MODEL = os.getenv("WHISPER_MODEL", "base")
WHISPER_DEVICE = os.getenv("WHISPER_DEVICE", "cpu")
BACKEND_API_URL = os.getenv("BACKEND_API_URL", "http://backend:8000/api")
OUTPUT_PATH = Path(os.getenv("OUTPUT_PATH", "/app/output"))
TEMPLATE_PATH = Path(os.getenv("DOCUMENT_TEMPLATE_PATH", "/app/templates"))
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO")

# Logging Setup
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=getattr(logging, LOG_LEVEL),
)
logger = logging.getLogger(__name__)


class VoiceAssistant:
    """Hauptklasse für den Telegram Voice Assistant"""

    def __init__(self):
        self.http_client = httpx.AsyncClient(timeout=30.0)
        OUTPUT_PATH.mkdir(parents=True, exist_ok=True)
        # Whisper Model laden (nur wenn lokal)
        self.whisper_model = None
        if not OPENAI_API_KEY:
            logger.info(f"Lade lokales Whisper-Modell: {WHISPER_MODEL} auf {WHISPER_DEVICE}")
            self.whisper_model = whisper.load_model(WHISPER_MODEL, device=WHISPER_DEVICE)

    async def transcribe_audio(self, audio_file_path: str) -> str:
        """
        Transkribiert Audio zu Text
        
        Args:
            audio_file_path: Pfad zur Audio-Datei
            
        Returns:
            Transkribierter Text
        """
        if OPENAI_API_KEY:
            # OpenAI Whisper API
            logger.info("Transkription via OpenAI Whisper API")
            with open(audio_file_path, "rb") as audio_file:
                response = await self.http_client.post(
                    "https://api.openai.com/v1/audio/transcriptions",
                    headers={"Authorization": f"Bearer {OPENAI_API_KEY}"},
                    files={"file": audio_file},
                    data={"model": "whisper-1", "language": "de"},
                )
                response.raise_for_status()
                return response.json()["text"]
        else:
            # Lokales Whisper-Modell
            logger.info("Transkription via lokales Whisper-Modell")
            result = self.whisper_model.transcribe(audio_file_path, language="de")
            return result["text"]

    async def generate_markdown(self, text: str, title: str) -> str:
        """
        Generiert Markdown-Dokument aus Text
        
        Args:
            text: Transkribierter Text
            title: Dokumenten-Titel
            
        Returns:
            Pfad zum generierten Markdown-Dokument
        """
        timestamp = datetime.now().isoformat()
        filename = f"{title.replace(' ', '_')}_{timestamp.replace(':', '-')}.md"
        filepath = OUTPUT_PATH / filename

        content = f"""# {title}

**Erstellt:** {timestamp}  
**Quelle:** Telegram Voice Message

---

## Transkription

{text}

---

*Generiert von NOVA v3 Telegram Voice Assistant*
"""
        filepath.write_text(content, encoding="utf-8")
        logger.info(f"Markdown generiert: {filepath}")
        return str(filepath)

    async def generate_pdf(self, markdown_path: str) -> str:
        """
        Konvertiert Markdown zu PDF
        
        Args:
            markdown_path: Pfad zum Markdown-Dokument
            
        Returns:
            Pfad zum generierten PDF
        """
        pdf_path = Path(markdown_path).with_suffix(".pdf")
        
        # Markdown zu HTML
        md_content = Path(markdown_path).read_text(encoding="utf-8")
        html_content = markdown(md_content, extras=["tables", "fenced-code-blocks"])
        
        # HTML zu PDF mit WeasyPrint
        HTML(string=html_content).write_pdf(pdf_path)
        logger.info(f"PDF generiert: {pdf_path}")
        return str(pdf_path)

    async def generate_docx(self, text: str, title: str) -> str:
        """
        Generiert DOCX-Dokument aus Text
        
        Args:
            text: Transkribierter Text
            title: Dokumenten-Titel
            
        Returns:
            Pfad zum generierten DOCX
        """
        timestamp = datetime.now().isoformat()
        filename = f"{title.replace(' ', '_')}_{timestamp.replace(':', '-')}.docx"
        filepath = OUTPUT_PATH / filename

        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Erstellt: {timestamp}")
        doc.add_paragraph("Quelle: Telegram Voice Message")
        doc.add_heading("Transkription", level=1)
        doc.add_paragraph(text)
        doc.add_paragraph("\nGeneriert von NOVA v3 Telegram Voice Assistant")
        
        doc.save(filepath)
        logger.info(f"DOCX generiert: {filepath}")
        return str(filepath)

    async def save_to_backend(self, text: str, title: str, file_paths: list):
        """
        Speichert Transkription im NOVA Backend
        
        Args:
            text: Transkribierter Text
            title: Dokumenten-Titel
            file_paths: Liste der generierten Dateipfade
        """
        try:
            response = await self.http_client.post(
                f"{BACKEND_API_URL}/transcriptions",
                json={
                    "title": title,
                    "text": text,
                    "source": "telegram_voice",
                    "files": [str(p) for p in file_paths],
                },
            )
            response.raise_for_status()
            logger.info("Transkription im Backend gespeichert")
        except Exception as e:
            logger.error(f"Fehler beim Speichern im Backend: {e}")


# Telegram Bot Handlers
assistant = VoiceAssistant()


async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler für /start Kommando"""
    await update.message.reply_text(
        "🎤 **NOVA v3 Voice Assistant**\n\n"
        "Sende mir eine Sprachnachricht und ich erstelle daraus ein Dokument!\n\n"
        "**Befehle:**\n"
        "/start - Diese Nachricht\n"
        "/help - Hilfe anzeigen\n"
        "/formats - Verfügbare Formate"
    )


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler für /help Kommando"""
    await update.message.reply_text(
        "**Verwendung:**\n\n"
        "1. Sende eine Sprachnachricht\n"
        "2. Ich transkribiere sie automatisch\n"
        "3. Du erhältst ein Dokument (Markdown, PDF, DOCX)\n\n"
        "**Tipps:**\n"
        "- Spreche deutlich und in ganzen Sätzen\n"
        "- Vermeide Hintergrundgeräusche\n"
        "- Maximal 10 Minuten pro Nachricht"
    )


async def formats_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler für /formats Kommando"""
    await update.message.reply_text(
        "**Verfügbare Formate:**\n\n"
        "📄 **Markdown (.md)** - Immer generiert\n"
        "📕 **PDF (.pdf)** - Automatisch aus Markdown\n"
        "📘 **DOCX (.docx)** - Microsoft Word Format\n\n"
        "Alle Formate werden automatisch erstellt und dir gesendet."
    )


async def voice_message_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler für Sprachnachrichten"""
    user = update.effective_user
    logger.info(f"Sprachnachricht empfangen von {user.username} ({user.id})")

    # Status-Nachricht
    status_msg = await update.message.reply_text("🎤 Verarbeite Sprachnachricht...")

    try:
        # Audio-Datei herunterladen
        voice = update.message.voice
        file = await context.bot.get_file(voice.file_id)
        
        with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as tmp_file:
            await file.download_to_drive(tmp_file.name)
            audio_path = tmp_file.name

        # Transkription
        await status_msg.edit_text("✍️ Transkribiere Audio...")
        text = await assistant.transcribe_audio(audio_path)
        
        if not text.strip():
            await status_msg.edit_text("❌ Keine Sprache erkannt. Bitte versuche es erneut.")
            return

        # Dokumente generieren
        await status_msg.edit_text("📝 Generiere Dokumente...")
        title = f"Transkription_{user.username}"
        
        md_path = await assistant.generate_markdown(text, title)
        pdf_path = await assistant.generate_pdf(md_path)
        docx_path = await assistant.generate_docx(text, title)

        # Backend-Speicherung
        await assistant.save_to_backend(text, title, [md_path, pdf_path, docx_path])

        # Dokumente senden
        await status_msg.edit_text("📤 Sende Dokumente...")
        
        await update.message.reply_text(f"✅ **Transkription abgeschlossen!**\n\n{text[:500]}...")
        
        with open(md_path, "rb") as f:
            await update.message.reply_document(document=f, filename=Path(md_path).name)
        
        with open(pdf_path, "rb") as f:
            await update.message.reply_document(document=f, filename=Path(pdf_path).name)
        
        with open(docx_path, "rb") as f:
            await update.message.reply_document(document=f, filename=Path(docx_path).name)

        await status_msg.delete()
        
        # Temporäre Dateien löschen
        os.unlink(audio_path)

    except Exception as e:
        logger.error(f"Fehler bei Verarbeitung: {e}", exc_info=True)
        await status_msg.edit_text(f"❌ Fehler: {str(e)}")


def main():
    """Hauptfunktion - Startet den Bot"""
    if not TELEGRAM_BOT_TOKEN:
        logger.error("TELEGRAM_BOT_TOKEN nicht gesetzt!")
        return

    logger.info("Starte NOVA v3 Telegram Voice Assistant...")

    # Application erstellen
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()

    # Handler registrieren
    application.add_handler(CommandHandler("start", start_command))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("formats", formats_command))
    application.add_handler(MessageHandler(filters.VOICE, voice_message_handler))

    # Bot starten
    logger.info("Bot läuft... (Ctrl+C zum Beenden)")
    application.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()